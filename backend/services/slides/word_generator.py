from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class TalkingScriptWordGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置Word文档样式"""
        # 标题样式
        title_style = self.doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_format = title_style.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_after = Pt(12)
        
        font = title_style.font
        font.name = 'Arial'
        font.size = Pt(16)
        font.bold = True
        
        # 幻灯片标题样式
        slide_title_style = self.doc.styles.add_style('Slide Title', WD_STYLE_TYPE.PARAGRAPH)
        slide_title_format = slide_title_style.paragraph_format
        slide_title_format.space_before = Pt(12)
        slide_title_format.space_after = Pt(6)
        
        slide_font = slide_title_style.font
        slide_font.name = 'Arial'
        slide_font.size = Pt(14)
        slide_font.bold = True
        slide_font.color.rgb = None  # 使用默认颜色
        
        # 正文样式
        body_style = self.doc.styles.add_style('Script Body', WD_STYLE_TYPE.PARAGRAPH)
        body_format = body_style.paragraph_format
        body_format.space_after = Pt(6)
        body_format.line_spacing = 1.5
        
        body_font = body_style.font
        body_font.name = 'Times New Roman'
        body_font.size = Pt(12)
        
        # 注释样式
        note_style = self.doc.styles.add_style('Speaking Notes', WD_STYLE_TYPE.PARAGRAPH)
        note_format = note_style.paragraph_format
        note_format.space_after = Pt(3)
        note_format.left_indent = Inches(0.5)
        
        note_font = note_style.font
        note_font.name = 'Arial'
        note_font.size = Pt(10)
        note_font.italic = True
    
    def generate_document(self, scripts_data, presentation_title="Presentation Talking Script"):
        """生成完整的Word文档"""
        # 添加文档标题
        title = self.doc.add_paragraph(presentation_title, style='Custom Title')
        
        # 添加文档信息
        info_para = self.doc.add_paragraph()
        info_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Total Slides: {len(scripts_data)}\n")
        if scripts_data:
            info_para.add_run(f"Script Style: {scripts_data[0].get('script_style', 'academic').title()}\n")
        
        total_duration = self._calculate_total_duration(scripts_data)
        info_para.add_run(f"Estimated Total Duration: {total_duration}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分页符
        self.doc.add_page_break()
        
        # 添加目录标题
        toc_title = self.doc.add_paragraph("Table of Contents", style='Custom Title')
        
        # 生成目录
        for script in scripts_data:
            toc_entry = self.doc.add_paragraph()
            toc_entry.add_run(f"Slide {script['slide_number']}: {script['slide_title']}")
            toc_entry.add_run(f" .................. Page {script['slide_number'] + 2}")
        
        # 添加分页符
        self.doc.add_page_break()
        
        # 为每个幻灯片生成演讲稿页面
        for i, script in enumerate(scripts_data):
            self._add_slide_script(script)
            
            # 除了最后一个幻灯片，都添加分页符
            if i < len(scripts_data) - 1:
                self.doc.add_page_break()
        
        return self.doc
    
    def _add_slide_script(self, script_data):
        """添加单个幻灯片的演讲稿"""
        # 幻灯片标题
        slide_title = f"Slide {script_data['slide_number']}: {script_data['slide_title']}"
        self.doc.add_paragraph(slide_title, style='Slide Title')
        
        # 幻灯片要点概览
        points_title = self.doc.add_paragraph("Key Points Overview:")
        points_title.style = self.doc.styles['Heading 3']
        
        for point in script_data['slide_content_points']:
            point_para = self.doc.add_paragraph()
            point_para.add_run(f"• {point}")
            point_para.paragraph_format.left_indent = Inches(0.25)
        
        # 演讲稿部分
        script_title = self.doc.add_paragraph("Talking Script:")
        script_title.style = self.doc.styles['Heading 3']
        
        talking_script = script_data['talking_script']
        
        # 介绍部分
        if talking_script.get('intro'):
            intro_header = self.doc.add_paragraph("Introduction:")
            intro_header.style = self.doc.styles['Heading 4']
            intro_para = self.doc.add_paragraph(talking_script['intro'], style='Script Body')
        
        # 主体部分
        if talking_script.get('main_body'):
            main_header = self.doc.add_paragraph("Main Content:")
            main_header.style = self.doc.styles['Heading 4']
            for paragraph in talking_script['main_body']:
                main_para = self.doc.add_paragraph(paragraph, style='Script Body')
        
        # 结论部分
        if talking_script.get('conclusion'):
            conclusion_header = self.doc.add_paragraph("Conclusion:")
            conclusion_header.style = self.doc.styles['Heading 4']
            conclusion_para = self.doc.add_paragraph(talking_script['conclusion'], style='Script Body')
        
        # 演讲提示
        notes_title = self.doc.add_paragraph("Speaking Notes:")
        notes_title.style = self.doc.styles['Heading 3']
        
        notes_para = self.doc.add_paragraph(style='Speaking Notes')
        notes_para.add_run(f"• Estimated Duration: {script_data['estimated_duration']}\n")
        notes_para.add_run(f"• Word Count: {script_data['word_count']} words\n")
        
        speaking_cues = script_data['speaking_cues']
        if speaking_cues['total_cues'] > 0:
            notes_para.add_run(f"• Speaking Cues: {speaking_cues['pauses']} pauses, ")
            notes_para.add_run(f"{speaking_cues['emphasis']} emphasis points, ")
            notes_para.add_run(f"{speaking_cues['slow_delivery']} slow delivery points\n")
        
        notes_para.add_run("• Remember to maintain eye contact with the audience\n")
        notes_para.add_run("• Speak clearly and at an appropriate pace\n")
        notes_para.add_run("• Use natural gestures to emphasize key points")
    
    def _calculate_total_duration(self, scripts_data):
        """计算总演讲时间"""
        total_words = sum(script.get('word_count', 0) for script in scripts_data)
        # 假设平均每分钟150个单词
        estimated_minutes = total_words / 150
        return f"approximately {estimated_minutes:.1f} minutes"
    
    def save_document(self, file_path):
        """保存文档到指定路径"""
        self.doc.save(file_path)
        return file_path

def generate_talking_script_word(scripts_data, output_path, presentation_title="Presentation Talking Script"):
    """生成talking script的Word文档的便捷函数"""
    generator = TalkingScriptWordGenerator()
    doc = generator.generate_document(scripts_data, presentation_title)
    return generator.save_document(output_path) 