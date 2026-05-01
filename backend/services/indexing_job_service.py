"""
Async indexing job service.

Manages background document indexing tasks with status tracking in MongoDB.
Uses asyncio.create_task for in-process background work (no Celery needed).
"""
from __future__ import annotations

import asyncio
import hashlib
import logging
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from backend.core.database import db
from backend.config import Config
from backend.services.file_asset_service import register_file_asset

logger = logging.getLogger(__name__)

COLLECTION = "indexing_jobs"


async def create_job(
    course_id: str,
    filename: str,
    content_bytes: bytes,
    user_id: str,
    chapter_id: str = "",
    use_fast_extract: bool = False,
) -> dict:
    """Create a new indexing job and start processing in the background.

    Returns the job document with id, status, etc.
    If an identical file (same course + filename + content hash) was already
    indexed, returns a fast "duplicate" response without re-indexing.
    """
    content_hash = hashlib.sha256(content_bytes).hexdigest()
    print(f"[DEBUG create_job] course={course_id} file={filename} hash={content_hash[:12]}")

    # Check for duplicate: same course, filename, content hash, already done
    existing = await db[COLLECTION].find_one({
        "course_id": course_id,
        "filename": filename,
        "content_hash": content_hash,
        "chapter_id": str(chapter_id or ""),
        "status": "done",
    }, sort=[("created_at", -1)])

    if existing:
        # Verify the document actually exists in the vectorstore meta.json.
        # If meta was wiped (e.g. prior failed delete), we must re-index.
        from backend.services.course_rag_service import course_rag_service
        indexed_names = {
            d["doc_name"]
            for d in course_rag_service.list_indexed_documents(course_id)
        }
        if filename in indexed_names:
            print(f"[DEBUG create_job] DUPLICATE confirmed in meta, returning existing job_id={existing['job_id']}")
            logger.info("Skipping duplicate index: course=%s file=%s hash=%s", course_id, filename, content_hash[:12])
            return {
                "job_id": existing["job_id"],
                "status": "done",
                "filename": filename,
                "content_hash": content_hash,
                "duplicate": True,
            }
        else:
            # MongoDB says done but meta.json disagrees — invalidate the stale record
            print(f"[DEBUG create_job] STALE duplicate: meta.json missing '{filename}', invalidating old job and re-indexing")
            await db[COLLECTION].update_one(
                {"_id": existing["_id"]},
                {"$set": {"status": "stale", "updated_at": datetime.now(timezone.utc)}},
            )

    job_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)

    course_dir = Path(Config.KNOWLEDGE_BASE_UPLOAD_DIR) / course_id
    course_dir.mkdir(parents=True, exist_ok=True)
    source_name = f"{job_id}_{filename}"
    source_rel_path = Path("uploads") / "knowledge_base" / course_id / source_name
    source_abs_path = Path(Config.BASE_DIR) / source_rel_path
    source_abs_path.write_bytes(content_bytes)

    job_doc = {
        "job_id": job_id,
        "course_id": course_id,
        "filename": filename,
        "content_hash": content_hash,
        "file_size": len(content_bytes),
        "user_id": user_id,
        "chapter_id": str(chapter_id or ""),
        "use_fast_extract": use_fast_extract,
        "status": "pending",  # pending -> processing -> done | failed
        "phase": "pending",
        "progress": 0,
        "error": None,
        "result": None,
        "created_at": now,
        "updated_at": now,
        "source_path": source_rel_path.as_posix(),
    }
    await db[COLLECTION].insert_one(job_doc)

    try:
        await register_file_asset(
            file_type="knowledge_source",
            storage_path=source_rel_path.as_posix(),
            size=len(content_bytes),
            owner_type="knowledge_document",
            owner_id=job_id,
            created_by=user_id,
            filename=filename,
            course_id=course_id,
            scope="knowledge",
            user_id=user_id,
            metadata={
                "job_id": job_id,
                "content_hash": content_hash,
                "chapter_id": str(chapter_id or ""),
            },
        )
    except Exception:
        logger.exception("Failed to register knowledge source file asset")

    # Fire-and-forget background task
    asyncio.create_task(_process_job(job_id, course_id, filename, chapter_id=str(chapter_id or ""), use_fast_extract=use_fast_extract))

    return {
        "job_id": job_id,
        "status": "pending",
        "filename": filename,
        "content_hash": content_hash,
    }


async def get_job_status(job_id: str) -> Optional[dict]:
    """Return the current status of an indexing job."""
    doc = await db[COLLECTION].find_one({"job_id": job_id}, {"_id": 0})
    if doc and "created_at" in doc:
        doc["created_at"] = doc["created_at"].isoformat()
    if doc and "updated_at" in doc:
        doc["updated_at"] = doc["updated_at"].isoformat()
    return doc


async def _update_status(job_id: str, status: str, **extra) -> None:
    update = {"status": status, "updated_at": datetime.now(timezone.utc)}
    update.update(extra)
    await db[COLLECTION].update_one({"job_id": job_id}, {"$set": update})


async def _update_progress(job_id: str, progress: int, phase: str = "") -> None:
    """Update the indexing progress (0-100) and optional phase label for a job."""
    fields: dict = {"progress": progress, "updated_at": datetime.now(timezone.utc)}
    if phase:
        fields["phase"] = phase
    await db[COLLECTION].update_one({"job_id": job_id}, {"$set": fields})


async def _process_job(
    job_id: str,
    course_id: str,
    filename: str,
    chapter_id: str = "",
    use_fast_extract: bool = False,
) -> None:
    """Background coroutine that extracts text and indexes it."""
    try:
        await _update_status(job_id, "processing")
        await _update_progress(job_id, 5, "extracting")
        print(f"\n[DEBUG _process_job] START job={job_id} course={course_id} file={filename} fast={use_fast_extract}")

        source_path = Path(Config.BASE_DIR) / "uploads" / "knowledge_base" / course_id / f"{job_id}_{filename}"
        print(f"[DEBUG _process_job] source_path={source_path} exists={source_path.exists()}")

        # Extract text (run blocking I/O in executor)
        text = await asyncio.get_event_loop().run_in_executor(
            None, _extract_text_from_path, source_path, use_fast_extract
        )

        print(f"[DEBUG _process_job] extracted text length={len(text) if text else 0}")
        if not text or not text.strip():
            print(f"[DEBUG _process_job] EMPTY TEXT - marking failed")
            await _update_status(job_id, "failed", error="Could not extract any text from the file")
            return

        # Index (also blocking — ChromaDB operations)
        # Progress 60-100% maps to indexing phase; extraction is 5-60%.
        await _update_progress(job_id, 60, "indexing")
        loop = asyncio.get_event_loop()

        def _progress_cb(pct: int):
            # Map index_document's 0-100 → overall 60-100
            mapped = 60 + int(pct * 0.4)
            asyncio.run_coroutine_threadsafe(_update_progress(job_id, mapped, "indexing"), loop)

        from backend.services.course_rag_service import course_rag_service
        print(f"[DEBUG _process_job] calling index_document course={course_id} doc={filename}")
        result = await asyncio.get_event_loop().run_in_executor(
            None, course_rag_service.index_document, course_id, filename, text, chapter_id, _progress_cb
        )
        print(f"[DEBUG _process_job] index_document returned: {result}")

        vectorstore_path = Path(Config.RAG_VECTORSTORE_DIR) / "courses" / course_id
        try:
            existing_vector_asset = await db.file_assets.find_one(
                {"file_type": "knowledge_vectorstore", "owner_type": "course", "owner_id": course_id}
            )
            if not existing_vector_asset:
                await register_file_asset(
                    file_type="knowledge_vectorstore",
                    storage_path=vectorstore_path.relative_to(Path(Config.BASE_DIR)).as_posix(),
                    size=0,
                    owner_type="course",
                    owner_id=course_id,
                    created_by="system",
                    filename=f"course_{course_id}_vectorstore",
                    course_id=course_id,
                    scope="knowledge",
                    metadata={"managed_by": "course_rag_service"},
                )
        except Exception:
            logger.exception("Failed to register vectorstore asset")

        await _update_status(job_id, "done", result=result)
        print(f"[DEBUG _process_job] DONE job={job_id}")
        logger.info("Indexing job %s completed: %s", job_id, result)

    except Exception as exc:
        print(f"[DEBUG _process_job] EXCEPTION job={job_id}: {exc}")
        import traceback; traceback.print_exc()
        logger.exception("Indexing job %s failed", job_id)
        await _update_status(job_id, "failed", error="Internal indexing error")


def _extract_text_from_path(source_path: Path, use_fast: bool = False) -> str:
    """Extract text from a persisted source file. Runs in thread executor."""
    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        from backend.utils.pdf_extractor import extract_text_from_pdf

        return extract_text_from_pdf(str(source_path), use_fast=use_fast)

    if suffix == ".docx":
        return _extract_text_from_docx(source_path)

    content = source_path.read_bytes().decode("utf-8", errors="replace")

    # Strip Markdown syntax for .md/.markdown files to avoid
    # noise like ## / ** / []() leaking into vector embeddings.
    if suffix in (".md", ".markdown"):
        content = _strip_markdown(content)

    return content


def _strip_markdown(text: str) -> str:
    """Remove common Markdown formatting while preserving heading text and structure."""
    import re
    # Keep heading text but strip # markers (preserve \n for section splitting)
    text = re.sub(r"^(#{1,6})\s+", "", text, flags=re.MULTILINE)
    # Bold / italic
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # Strikethrough
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    # Inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Links [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Images ![alt](url) → alt
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # List markers (preserve text)
    text = re.sub(r"^(\s*)[*\-+]\s+", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^(\s*)\d+\.\s+", r"\1", text, flags=re.MULTILINE)
    # Blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    return text


def _extract_text_from_docx(source_path: Path) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx not installed; cannot process %s", source_path.name)
        return source_path.read_bytes().decode("utf-8", errors="replace")

    try:
        doc = DocxDocument(str(source_path))
    except Exception:
        logger.exception("Failed to open DOCX: %s", source_path.name)
        return ""

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Map DOCX heading styles to Markdown-style headings so the
        # section parser can detect them downstream.
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
            except ValueError:
                level = 1
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # Extract tables as Markdown-formatted tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # Insert header separator after first row
            header_sep = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
            rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    return "\n\n".join(parts)
